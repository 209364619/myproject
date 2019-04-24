# coding=utf-8
from docx import Document


def word_test():
    document = Document()

    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
    document.add_heading('The REAL meaning of the universe')
    document.add_heading('The REAL meaning of the universe')
    document.add_page_break()
    table = document.add_table(rows=2, cols=2)
    cell = table.cell(0, 1)
    cell = table.cell(0, 1)
    row = table.rows[1]
    row.cells[0].text = 'Foo bar to you.'
    row.cells[1].text = 'And a hearty foo bar to you too sir!'

    document.add_picture('image-filename.png')

    document.save('test.docx')


def read_paragraph(filepath):
    document = Document(filepath)
    for paragraph in document.paragraphs:
        print(paragraph.text)


def read_table(filepath):
    document = Document(filepath)
    tables = document.tables
    print len(tables)
    for i in range(len(tables)):
        # 遍历所有表格
        print 'table' + str(i)
        table = tables[i]
        for m in range(len(table.rows)):
            for n in range(len(table.columns)):
                print table.cell(m, n).text + '\t',
            print


if __name__ == '__main__':
    read_table(filepath='test.docx')
